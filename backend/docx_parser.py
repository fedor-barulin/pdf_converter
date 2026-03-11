import os
import win32com.client
from docx import Document
from logger import get_logger

logger = get_logger()

def convert_doc_to_docx(doc_path: str) -> str:
    """Конвертирует .doc файл в .docx с использованием win32com (должен быть установлен Word)"""
    docx_path = doc_path + 'x'
    if os.path.exists(docx_path):
        return docx_path
    
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs(os.path.abspath(docx_path), FileFormat=16) # 16 = wdFormatXMLDocument
        doc.Close()
        word.Quit()
        logger.info(f"Сконвертирован DOC в DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        logger.error(f"Ошибка конвертации DOC в DOCX: {e}")
        raise e

def extract_docx(file_path: str) -> list:
    """Извлекает текст и структуру из DOCX файла."""
    if file_path.lower().endswith('.doc'):
        file_path = convert_doc_to_docx(file_path)
        
    doc = Document(file_path)
    pages_data = []
    
    # В DOCX нет понятия страниц как таковых, поэтому эмулируем одну длинную "страницу"
    # Но мы можем извлекать параграфы и таблицы.
    lines = []
    tables_data = []
    
    # Извлечение параграфов (с учетом стилей для заголовков)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        
        is_heading = p.style.name.startswith("Heading") or "Заголовок" in p.style.name
        lines.append({
            "text": text,
            "is_heading": is_heading,
            "style": p.style.name
        })

    # Извлечение таблиц
    for idx, table in enumerate(doc.tables):
        table_rows = []
        header = []
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if r_idx == 0:
                header = row_data
            else:
                table_rows.append(row_data)
        
        tables_data.append({
            "table_id": f"table_{idx+1}",
            "header": header,
            "rows": table_rows
        })
        
    pages_data.append({
        "page": 1,
        "lines": lines,
        "tables": tables_data,
        "source": os.path.basename(file_path)
    })
    
    return pages_data
